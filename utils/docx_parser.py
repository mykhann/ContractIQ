import re
import logging
from typing import Optional
from io import BytesIO

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_bytes: bytes) -> Optional[str]:
    try:
        from docx import Document
        doc = Document(BytesIO(file_bytes))
        parts = []

        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(f"[TABLE] {row_text}")

        if not parts:
            return None

        full = "\n\n".join(parts)
        full = re.sub(r"\n{3,}", "\n\n", full)
        logger.info(f"DOCX: {len(full)} chars, {len(parts)} paragraphs")
        return full

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Could not parse DOCX: {e}")
